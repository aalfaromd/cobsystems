import streamlit as st
from docx import Document
from datetime import datetime
from io import BytesIO


st.title("📄 DAN Subrogation/Coordination of Benefits Cover Letter Generator")

st.markdown("Fill in the fields below to generate a professionally formatted subrogation letter as a downloadable Word document.")

# Collect user inputs
insurance_company = st.text_input("Insurance Company")
address_line1 = st.text_input("Address Line 1")
address_line2 = st.text_input("Address Line 2")
patient_name = st.text_input("Patient/Insured Name")
policy_number = st.text_input("DAN Policy/Member Number")
dob = st.date_input("Patient’s Date of Birth")
coverage_type = st.text_input("Type of Coverage")
dates_of_service = st.text_input("Date(s) of Service")
PrimaryIns_id = st.text_input("Primary Insurance ID")
provider_payment_details = st.text_input("Provider Payment Details")
include_claim_form = st.checkbox("Include completed claim form line")

if st.button("Generate Letter"):
    doc = Document()

    doc.add_paragraph("DAN Insurance Company\n\nSERVICES\n\n")
    doc.add_picture("dan logo letter.png", width=Inches(1.5))
    doc.add_paragraph(datetime.today().strftime("%B %d, %Y"))
    doc.add_paragraph(f"{insurance_company}\nAttn: Coordination of Benefits\n{address_line1}\n{address_line2}")
    doc.add_paragraph("\nRE: Coordination of Benefits Request - Out-of-Country Emergency Medical Expense\n")
    doc.add_paragraph(f"Patient/Insured Name: {patient_name}")
    doc.add_paragraph(f"DAN Policy/Member Number: {policy_number}")
    doc.add_paragraph(f"Patient’s DOB: {dob.strftime('%m/%d/%Y')}")
    doc.add_paragraph(f"Type of Coverage: {coverage_type}")
    doc.add_paragraph(f"Date(s) of Service: {dates_of_service}")
    doc.add_paragraph(f"Primary Insurance ID: {PrimaryIns_id}")
    doc.add_paragraph(
        "\nDear Sir/Madam,\n\n"
        "We have recently completed a review of claim expenses incurred and paid by our company that relate to the same loss "
        "for which the insured’s policy with your company also covers.\n\n"
        "As the coverage under our policy is Excess (secondary) only coverage, we respectfully submit the following documents for your consideration:\n"
        "· A copy of our claim with pertinent coverage documents, each medical invoice, and COB claim itemization. "
        "(Please note that EOBs for payment made by bank wire may not be available.)\n"
        f"· Proof of payment made to {provider_payment_details}."
    )

    if include_claim_form:
        doc.add_paragraph("· A copy of your insured’s completed claim form assigning benefits to our company.")

    doc.add_paragraph(
        "\nPlease determine the share of expenses payable under your policy and remit payment by check to:\n\n"
        "DAN Services Inc.                     (Tax ID # 56-2089710)\n"
        "Attn: DAN Claims Recoveries\n"
        "6 West Colony Place\n"
        "Durham, NC  27705\n"
    )

    doc.add_paragraph(
        "Your cooperation in this matter is greatly appreciated. "
        "If additional information is needed, please contact the undersigned at the contact below "
        "or direct by phone at 919-684-2948 ext. 1222 or email at aalfaro@dan.org."
    )

    doc.add_paragraph(
        "Sincerely,\n\nAngelica Alfaro\nRecovery Team\n\nEnclosures\n\n"
        "DAN Claims ( 6 W Colony Place ( Durham, NC  ( 27705 ( USA\n"
        "Phone 1.919.493.0912 ( Fax 1.913.493.3040 ( Claims@dan.org\n"
        "DAN Services, Inc. is a for-profit Subsidiary of Divers Alert Network"
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("✅ Letter generated!")
    st.download_button(
        label="📥 Download Subrogation Letter",
        data=buffer,
        file_name=f"Subrogation_Letter_{patient_name.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
